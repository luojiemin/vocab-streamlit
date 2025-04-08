import streamlit as st
import pandas as pd
import easyocr
from PIL import Image
import io
from docx import Document

# 初始化 OCR 引擎
reader = easyocr.Reader(['en', 'ch_sim'])

# 配置页面
st.set_page_config(page_title="英语词汇扩展系统", layout="centered")
st.title("📘 高三常忘英语词汇扩展学习系统")

# 上传截图
uploaded_files = st.file_uploader("上传中英文截图（支持多张）", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

# 初始化表格
results = []

# 示例词汇补全函数（用于真实部署时替换为完整词库逻辑）
def enrich_word_data(word, meaning):
    sample_dict = {
        "violate": {
            "pos": "v.",
            "collocation": "violate the law（违法）",
            "example": "He was fined for violating traffic rules.（他因违反交通规则被罚款。）",
            "derivatives": "violates (第三人称单数), violating (现在分词), violated (过去式/过去分词), violation (n. 违反), violator (n. 违规者)",
            "confusing": "violet（紫罗兰）violent（暴力的）"
        }
        # 更多词条可扩展...
    }
    data = sample_dict.get(word.lower(), {})
    return {
        "word": word,
        "pos": data.get("pos", ""),
        "cn_meaning": meaning,
        "collocation": data.get("collocation", ""),
        "example": data.get("example", ""),
        "derivatives": data.get("derivatives", ""),
        "confusing": data.get("confusing", "")
    }

# 处理上传图片
if uploaded_files:
    for file in uploaded_files:
        image = Image.open(file)
    import numpy as np
result = reader.readtext(np.array(image))

        text = "\n".join([line[1] for line in result])
        lines = text.strip().split('\n')
        for line in lines:
            parts = line.strip().split()
            if len(parts) >= 3 and parts[0].isdigit():
                word = parts[1]
                meaning = ''.join(parts[2:])
                result = enrich_word_data(word, meaning)
                results.append(result)

# 展示并导出结果
if results:
    df = pd.DataFrame(results)
    st.success(f"共提取到 {len(results)} 个词汇")
    st.dataframe(df)

    # 导出为 Word 文档
    doc = Document()
    doc.add_heading("高三英语常忘词扩展记忆手册", level=1)
    for item in results:
        doc.add_paragraph(f"■ {item['word']}  {item['pos']}  {item['cn_meaning']}")
        doc.add_paragraph(f"常见搭配：{item['collocation']}")
        doc.add_paragraph(f"例句：{item['example']}")
        doc.add_paragraph(f"词形变化：{item['derivatives']}")
        doc.add_paragraph(f"形近词辨析：{item['confusing']}")
        doc.add_paragraph("")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button("📄 下载 Word 文档", buffer, file_name="词汇扩展结果.docx")
else:
    st.info("请上传包含中英文单词的截图，系统将自动识别并补全词汇信息。")
